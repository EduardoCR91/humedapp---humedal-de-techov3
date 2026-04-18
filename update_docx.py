import docx

doc = docx.Document('Documentacion/Proyecto Grado EcoVigia.docx')

def replace_para(para, new_text):
    para.clear()
    para.add_run(new_text)

# 1. Update Síntesis
for p in doc.paragraphs:
    if "Este proyecto consiste en el diseño y" in p.text and "84,5 %" in p.text:
        replace_para(p, "En este proyecto diseñamos y desarrollamos una aplicación móvil para ayudar a mitigar el rápido deterioro que sufren los humedales de la localidad de Kennedy, en especial Techo y El Burro. Nuestra idea nació al ver cómo se ha perdido gran parte de estos ecosistemas por la urbanización y la falta de herramientas que nos permitan a los ciudadanos vigilar y cuidar nuestro entorno. Con esta app, buscamos ir más allá de la información básica y ofrecer una plataforma en la que la comunidad pueda participar activamente. Añadimos un módulo para que cualquier persona pueda reportar incidentes con fotos y ubicación exacta, integramos un chatbot interactivo para resolver dudas al instante y agregamos una sección para ver los datos de los sensores (como calidad de agua o aire) en tiempo real. Nuestro propósito es que la aplicación se convierta en una herramienta cotidiana y de uso real para proteger nuestros humedales.")
        break

# 2. Add Justification
found_objs = -1
for i, p in enumerate(doc.paragraphs):
    if "Objetivos del proyecto" in p.text:
        found_objs = i
        break

if found_objs != -1:
    para_obj = doc.paragraphs[found_objs]
    new_heading = para_obj.insert_paragraph_before('Justificación del Proyecto')
    # Try to copy style from something else if possible, or just use bold
    try:
        new_heading.style = 'Heading 2'
    except:
        pass
    new_body = para_obj.insert_paragraph_before('El desarrollo de este proyecto se llevó a cabo en estrecha colaboración con el área de proyección social de la universidad, quienes nos guiaron a la hora de identificar las problemáticas reales de la comunidad y del ecosistema. Asimismo, contamos con el valioso acompañamiento del docente Ramiro Osorio. Él nos brindó asesoría técnica fundamental y nos ayudó directamente en la construcción del chatbot inteligente y en la estructuración del módulo de telemetría e IoT. Gracias a este trabajo conjunto, logramos enfocar la herramienta hacia una necesidad real del territorio, garantizando no solo su viabilidad técnica sino también su profundo impacto social.')

# 3. Add Point 5 Info (Entrega)
found_p5 = -1
for i, p in enumerate(doc.paragraphs):
    if "5." in p.text and "Entregable" in p.text: # just in case
        pass
    if "Diseño del software" in p.text and "ISO" in p.text:
        found_p5 = i
        break

if found_p5 != -1:
    para_diseño = doc.paragraphs[found_p5]
    nuevo = para_diseño.insert_paragraph_before('5. Entregables y Propiedades del Proyecto')
    try:
         nuevo.style = 'Heading 2'
    except:
         pass
    cuerpo = para_diseño.insert_paragraph_before('Cabe destacar que nuestro desarrollo tiene dos enfoques de entrega muy claros. A los usuarios finales (la comunidad, los visitantes y los colectivos ambientales), les entregaremos la aplicación móvil funcional y lista para usar. Por otro lado, y como aspecto central, nuestro cliente principal es la universidad. A la institución le entregaremos absolutamente todo el código fuente del proyecto, las arquitecturas y las bases de datos para que quede como la dueña del proyecto y del ecosistema digital. Esto garantiza que, a futuro, la universidad pueda escalar la aplicación de forma autónoma, conectar nuevos sensores o ampliar el contenido sin depender de terceros.')

# Humanize Metodología
for p in doc.paragraphs:
    if "El éxito de un proyecto de software no depende únicamente de la habilidad" in p.text:
        replace_para(p, "Para guiarnos en el proceso de creación de la aplicación, decidimos alejarnos de los métodos tradicionales y optamos por la metodología ágil bajo el marco de trabajo Scrum. Esto nos pareció la mejor opción porque nos permitió avanzar paso a paso, ir adaptando nuevas funciones sobre la marcha y así entregar valor constante a la comunidad en lugar de esperar meses para tener algo funcional.")
    if "Históricamente, el desarrollo de software se basaba en modelos predictivos" in p.text:
        replace_para(p, "En vez de seguir etapas rígidas e inamovibles, implementamos un desarrollo iterativo. Esto significa que dividimos el trabajo en pequeños ciclos ('Sprints'), lo que nos dio la flexibilidad de ir haciendo pruebas rápidas y corregir errores a tiempo.")
    if "La aplicación de gestión de humedales no sea un producto rígido" in p.text:
        replace_para(p, "Trabajar así nos asegura que la app no sea un bloque cerrado, sino que pueda crecer. Al separar todo en piezas (monitoreo, participación, el chat), cualquier estudiante de la universidad en el futuro puede agregar más sensores o cambiar cosas sin miedo a dañar lo que ya está hecho.")

doc.save('Proyecto de grado ecovigia.docx')
doc.save('Documentacion/Proyecto de grado ecovigia.docx')
