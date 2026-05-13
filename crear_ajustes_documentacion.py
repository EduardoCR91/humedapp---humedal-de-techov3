from docx import Document
from docx.shared import Pt
import os

doc = Document()

# Título Principal
doc.add_heading('Ajustes y Complementos para el Proyecto de Grado EcoVigía', level=1)
doc.add_paragraph('A continuación, se presentan los textos redactados para completar o actualizar las secciones del documento final de tesis, basándose en el estado actual del proyecto.')

# Sección 5.1
doc.add_heading('Para la Sección: 5.1 Herramientas utilizadas en el desarrollo', level=2)
doc.add_paragraph('Añadir las siguientes herramientas a la lista existente:')

doc.add_heading('@capacitor/assets y @capacitor/splash-screen', level=3)
doc.add_paragraph('Se implementó el conjunto de herramientas de Capacitor para la generación automatizada de recursos visuales nativos. @capacitor/assets permite, a partir de una única imagen maestra (icono y pantalla de carga), compilar automáticamente todos los tamaños y resoluciones requeridos para el ecosistema móvil de Android, optimizando el tiempo de empaquetado. Asimismo, @capacitor/splash-screen controla la duración, comportamiento y transiciones de la pantalla de carga inicial, asegurando que la aplicación no muestre destellos blancos mientras el motor de WebView (React) carga en segundo plano.')

doc.add_heading('Google Forms (Para pruebas UX)', level=3)
doc.add_paragraph('Plataforma utilizada para la estructuración y recolección de métricas de usabilidad. Facilitó la aplicación del cuestionario basado en la escala de Likert a los usuarios de prueba, permitiendo la generación automática de gráficos estadísticos sobre el nivel de satisfacción, facilidad de navegación y tiempo de respuesta de la aplicación.')

# Sección 6.3
doc.add_heading('Para la Sección: 6.3 Modificaciones realizadas', level=2)
doc.add_paragraph('Añadir este nuevo registro a la tabla de modificaciones, correspondiente a los ajustes de contraste realizados en el Módulo de Educación:')

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Table Grid'
h3 = table3.rows[0].cells
headers3 = ['ID', 'Hallazgo / Problema detectado', 'Solución Implementada', 'Fecha de ajuste']
for i, h in enumerate(headers3):
    h3[i].text = h

data3 = [
    ('M-04', 'Contraste deficiente en el Módulo de Educación: Los textos de entrada de fecha, hora y descripción en los formularios heredaban la regla global CSS de las tarjetas ecológicas, mostrándose de color verde brillante sobre fondo claro.', 'Se inyectaron estilos en línea directos (`color: "black"`) en las etiquetas de los componentes de React (`Education.tsx`) para anular la herencia de las reglas `!important` globales, restaurando el color oscuro y garantizando la legibilidad para el administrador y los usuarios.', 'Mayo de 2026')
]

for item in data3:
    row_cells = table3.add_row().cells
    for i in range(4):
        row_cells[i].text = item[i]

# Manual de Usuario
doc.add_heading('Para el: Manual de usuario / Guía de Configuración', level=2)
doc.add_paragraph('Añadir este bloque técnico a la guía de instalación y compilación del proyecto:')

doc.add_heading('Gestión de Íconos y Pantalla de Carga (Splash Screen) para Android', level=3)
doc.add_paragraph('Para modificar la identidad gráfica de la aplicación móvil antes de generar un nuevo APK, el sistema utiliza un enfoque de "fuente única" gestionado por Capacitor:')
doc.add_paragraph('1. Creación de recursos: El administrador o desarrollador debe ubicar dos archivos en la carpeta raíz `assets/`: \n- `icon.png` (Mínimo 1024x1024 píxeles, formato cuadrado).\n- `splash.png` (Mínimo 2732x2732 píxeles, formato cuadrado).')
doc.add_paragraph('2. Generación automática: Desde la terminal, se ejecuta el comando `npx @capacitor/assets generate`. Este script procesa las imágenes maestras y las recorta a todas las dimensiones estandarizadas requeridas por la carpeta nativa de Android.')
doc.add_paragraph('3. Configuración de duración: En el archivo `capacitor.config.ts`, la sección `SplashScreen` permite ajustar el tiempo de visualización (ej. `launchShowDuration: 4000` para 4 segundos) y el color de fondo para empalmar con la paleta de colores de la app.')

# Nueva Sección: Integración de Telemetría
doc.add_heading('Para la Sección: Arquitectura o Módulos del Sistema (Integración IoT)', level=2)
doc.add_paragraph('Añadir este bloque para documentar la integración interdisciplinaria con la carrera de Ingeniería de Telecomunicaciones para la ingesta de datos ambientales:')

doc.add_heading('Arquitectura de Ingesta de Datos Ambientales (Node-RED y Supabase Edge Functions)', level=3)
doc.add_paragraph('Para el componente de Monitoreo Ambiental en tiempo real, se estableció un flujo de datos (pipeline) colaborativo con el área de Ingeniería de Telecomunicaciones. Los sensores físicos desplegados en el humedal recolectan las variables críticas de calidad del agua y el ambiente, las cuales son procesadas localmente mediante la plataforma Node-RED.')
doc.add_paragraph('Para lograr la interoperabilidad con la aplicación móvil EcoVigía, se diseñó e implementó una API RESTful (Serverless) utilizando Supabase Edge Functions. Esta función actúa como un webhook seguro que recibe cargas útiles (payloads) en formato JSON desde Node-RED.')
doc.add_paragraph('Ejemplo de Payload recibido desde los sensores:')
doc.add_paragraph('{"temperature": 24.69, "ph": 17.82, "tds": 365, "ntu": 684.2, "quality": "Deficiente"}')
doc.add_paragraph('Estos datos se insertan automáticamente en la tabla relacional `climate_readings` de Supabase, permitiendo que la aplicación cliente construida en React consulte y renderice los indicadores ambientales históricos y en tiempo real a través del dashboard de Monitoreo, cerrando así la brecha entre el hardware físico (IoT) y el software de usuario final.')

# Save Document
doc.save('Documentacion/Ajustes_Finales_EcoVigia.docx')
print("Documento de ajustes generado correctamente en Documentacion/Ajustes_Finales_EcoVigia.docx")
